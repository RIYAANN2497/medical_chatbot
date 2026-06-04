import fitz
import base64
import os
import uuid
import shutil
from pathlib import Path
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from groq import Groq
import docx

CHROMA_BASE_DIR = str(Path(__file__).parent / "chroma_store")
SUPPORTED_IMAGES = {".jpg", ".jpeg", ".png", ".webp"}

_embeddings = None


def get_embeddings() -> HuggingFaceEmbeddings:
    global _embeddings
    if _embeddings is None:
        _embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={"device": "cpu"},
            encode_kwargs={"normalize_embeddings": True},
        )
    return _embeddings


def summarise_document(chunks: list[str], doc_name: str) -> str:
    from langchain_groq import ChatGroq
    from langchain_core.output_parsers import StrOutputParser
    if len(chunks) <= 6:
        selected = chunks
    else:
        selected = chunks[:3] + chunks[-3:]
    sample = "\n\n".join(selected)[:3000]

    prompt = f"""You are MediChat, a warm friendly medical buddy talking to a patient.

The patient just uploaded their report and wants to understand it. Talk to them like a caring friend explaining over coffee — casual, warm, plain English. No bullet points, no headers, no tables, no bold text, no formatting at all. Just natural flowing paragraphs like a real conversation.

Keep it under 150 words. Mention what the report is about, the key findings in simple words, anything that needs attention, and end with a warm note.

Document: {doc_name}

Report content:
{sample}

Explanation (plain conversational paragraphs only, no formatting):"""

    llm = ChatGroq(model_name="llama-3.3-70b-versatile", temperature=0.3)
    return (llm | StrOutputParser()).invoke(prompt)


# ── CHANGED: New image summariser specifically for medical images ──────────────
def summarise_image_document(raw_vision_text: str, doc_name: str) -> str:
    """
    Takes the raw clinical description from the vision model and produces
    a warm patient-friendly summary. Stored in summaries[] for the Docs tab.
    The raw_vision_text itself is stored in image_texts[] for ImageExplainerAgent.
    """
    from langchain_groq import ChatGroq
    from langchain_core.output_parsers import StrOutputParser

    prompt = f"""You are MediChat, a warm friendly medical buddy.

A patient uploaded a medical image called "{doc_name}". Below is a clinical description of what the image shows.

Write 2-3 warm, plain-English sentences summarising what this image is and what it generally shows.
No jargon, no bullet points, no formatting. Just a casual friendly sentence like you'd say over coffee.
Do NOT give a diagnosis. End with "Your doctor can explain exactly what this means for you."

Clinical description:
{raw_vision_text[:2000]}

Friendly summary (2-3 sentences only):"""

    llm = ChatGroq(model_name="llama-3.3-70b-versatile", temperature=0.3)
    return (llm | StrOutputParser()).invoke(prompt)


def extract_text_from_pdf(pdf_path: str, max_vision_pages: int = 10) -> str:
    with fitz.open(pdf_path) as doc:
        text = "".join(page.get_text() for page in doc)

    if len(text.strip()) >= 200:
        return text

    all_text: list[str] = []
    with fitz.open(pdf_path) as doc:
        pages_to_process = min(len(doc), max_vision_pages)
        for i in range(pages_to_process):
            img_path = f"{pdf_path}_page{i}.png"
            try:
                pix = doc[i].get_pixmap(dpi=200)
                pix.save(img_path)
                page_text = extract_text_from_image(img_path)
                if page_text.strip():
                    all_text.append(f"[Page {i + 1}]\n{page_text}")
            finally:
                if os.path.exists(img_path):
                    os.unlink(img_path)

    return "\n\n".join(all_text)


def extract_text_from_docx(docx_path: str) -> str:
    doc = docx.Document(docx_path)
    parts = [para.text for para in doc.paragraphs if para.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            seen: set[str] = set()
            cells: list[str] = []
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in seen:
                    seen.add(text)
                    cells.append(text)
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_text_from_image(image_path: str) -> str:
    """
    CHANGED: Prompt now auto-detects organ/body part and describes findings
    clinically for ANY radiological image (X-ray, MRI, CT, ultrasound, etc.).
    Also handles text-based images (lab reports, prescriptions) as before.
    """
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        raise EnvironmentError("GROQ_API_KEY is not set — cannot process image files.")

    client = Groq(api_key=api_key)
    ext = Path(image_path).suffix.lower()
    media_map = {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".webp": "image/webp",
    }
    media_type = media_map.get(ext, "image/jpeg")

    with open(image_path, "rb") as f:
        image_data = base64.standard_b64encode(f.read()).decode("utf-8")

    try:
        response = client.chat.completions.create(
            model="meta-llama/llama-4-scout-17b-16e-instruct",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{media_type};base64,{image_data}"
                            },
                        },
                        {
                            "type": "text",
                            # ── CHANGED: organ-agnostic clinical describe prompt ──
                            "text": (
                                "This is a medical image. First identify what type it is "
                                "and what body part or organ is shown.\n\n"
                                "Then based on what you see, describe ALL of the following "
                                "that are relevant:\n\n"
                                "WHAT IT IS:\n"
                                "- Type of scan (X-ray, MRI, CT, ultrasound, etc.)\n"
                                "- Body region and organ system shown (chest, abdomen, brain, "
                                "knee, spine, hand, pelvis, etc.)\n"
                                "- View or orientation if visible (PA, lateral, axial, "
                                "coronal, sagittal, etc.)\n\n"
                                "WHAT YOU SEE (describe only what is present in this image):\n"
                                "- Size, shape, and position of the primary organ(s) visible\n"
                                "- Surrounding structures and how they appear\n"
                                "- Any areas of abnormal density, opacity, shadow, or brightness\n"
                                "- Any fractures, breaks, dislocations, or bone abnormalities\n"
                                "- Any masses, nodules, fluid collections, swelling, or effusion\n"
                                "- Any asymmetry, displacement, or structural irregularity\n"
                                "- Any implants, foreign bodies, or medical devices visible\n"
                                "- Overall tissue and bone density and distribution\n\n"
                                "OVERALL IMPRESSION:\n"
                                "- What appears normal\n"
                                "- What appears abnormal or worth noting\n"
                                "- Clinical impression in one or two sentences\n\n"
                                "If this is NOT a radiological image but a text-based document "
                                "(lab report, prescription, discharge summary, handwritten notes):\n"
                                "- Extract ALL text, numbers, values, units, and findings "
                                "exactly as written.\n\n"
                                "Be thorough, specific, and clinical. Only describe what is "
                                "clearly visible in the image."
                            ),
                        },
                    ],
                }
            ],
            max_tokens=1024,
        )
        return response.choices[0].message.content
    except Exception as e:
        raise RuntimeError(f"Image extraction failed for {image_path}: {e}") from e


def ingest_multiple_files(
    file_paths: list,
    original_names: list,
    session_id: str = None,
) -> tuple:
    """
    CHANGED: Returns (vectorstore, chroma_dir, summaries, image_texts).
    image_texts holds the raw clinical vision descriptions for ImageExplainerAgent.
    summaries holds friendly patient-facing summaries for the Docs tab.

    NOTE: app.py must be updated to unpack 4 values from this return.
    """
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=80)

    all_chunks: list[Document] = []
    summaries: dict[str, str] = {}
    image_texts: dict[str, str] = {}   # ── CHANGED: raw clinical descriptions
    skipped: list[str] = []

    for file_path, original_name in zip(file_paths, original_names):
        ext = Path(original_name).suffix.lower()
        print(f"[ingest] Processing: {original_name}")

        try:
            if ext == ".pdf":
                raw_text = extract_text_from_pdf(file_path)
            elif ext == ".docx":
                raw_text = extract_text_from_docx(file_path)
            elif ext in SUPPORTED_IMAGES:
                raw_text = extract_text_from_image(file_path)
            elif ext == ".txt":
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    raw_text = f.read()
            else:
                print(f"[ingest] Skipping unsupported file: {original_name}")
                skipped.append(original_name)
                continue
        except Exception as e:
            print(f"[ingest] Error processing {original_name}: {e}")
            raise RuntimeError(f"Failed to process '{original_name}': {e}") from e

        chunks = splitter.split_text(raw_text)

        documents = [
            Document(
                page_content=chunk,
                metadata={
                    "source": original_name,
                    "chunk_index": i,
                    "file_type": ext.lstrip("."),
                },
            )
            for i, chunk in enumerate(chunks)
        ]
        all_chunks.extend(documents)

        # ── CHANGED: images get special handling ──────────────────────────────
        if ext in SUPPORTED_IMAGES:
            # Store raw clinical description for ImageExplainerAgent
            image_texts[original_name] = raw_text[:3000]
            # Store a short friendly summary for the Docs tab
            try:
                summaries[original_name] = summarise_image_document(raw_text, original_name)
            except Exception as e:
                print(f"[ingest] Image summary failed for {original_name}: {e}")
                summaries[original_name] = (
                    f"Medical image uploaded: {original_name}. "
                    "Your doctor can explain what this image shows."
                )
        else:
            # Text-based documents get the conversational summary as before
            try:
                summaries[original_name] = summarise_document(chunks, original_name)
            except Exception as e:
                print(f"[ingest] Summary failed for {original_name}: {e}")
                summaries[original_name] = "Summary unavailable for this document."
        # ── END CHANGED ───────────────────────────────────────────────────────

    if skipped:
        print(f"[ingest] Skipped unsupported files: {', '.join(skipped)}")

    if not all_chunks:
        raise ValueError(
            "No content could be extracted from the uploaded files. "
            "Check that your files are not empty or in an unsupported format."
        )

    embedding_model = get_embeddings()
    print(f"[ingest] Embedding {len(all_chunks)} total chunks (in-memory)")

    vectorstore = Chroma.from_documents(
        documents=all_chunks,
        embedding=embedding_model,
    )
    print("[ingest] Done.")
    # ── CHANGED: now returns 4 values instead of 3 ───────────────────────────
    return vectorstore, None, summaries, image_texts